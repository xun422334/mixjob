import os
import io
import json
from fastapi import APIRouter, UploadFile, File, Depends, HTTPException
from fastapi.responses import StreamingResponse
from sqlalchemy.orm import Session
from docx import Document
from docx.shared import Pt

from ..config import UPLOAD_DIR
from ..models.database import get_db, Resume, JobListing
from ..services.resume_parser import extract_text
from ..services.ai_service import structure_resume, update_resume

router = APIRouter(prefix="/api/resume", tags=["resume"])

ALLOWED_EXTENSIONS = {'.pdf', '.doc', '.docx', '.txt'}


@router.post("/upload")
async def upload_resume(file: UploadFile = File(...), db: Session = Depends(get_db)):
    # Validate file extension
    _, ext = os.path.splitext(file.filename or "")
    if ext.lower() not in ALLOWED_EXTENSIONS:
        raise HTTPException(
            status_code=400,
            detail=f"不支持的文件格式：{ext}。仅支持 PDF、DOC、DOCX、TXT",
        )

    os.makedirs(UPLOAD_DIR, exist_ok=True)

    file_path = os.path.join(UPLOAD_DIR, file.filename)
    content = await file.read()
    with open(file_path, "wb") as f:
        f.write(content)

    raw_text = extract_text(file_path, file.filename)

    try:
        structured = await structure_resume(raw_text)
    except Exception:
        structured = {"skills": [], "work_experience": [], "project_experience": [], "education": [], "personal_info": {}}

    resume = Resume(
        filename=file.filename,
        original_filename=file.filename,
        raw_text=raw_text,
        skills=structured.get("skills", []),
        experience=structured.get("work_experience", []),
        projects=structured.get("project_experience", []),
        education=structured.get("education", []),
        personal_info=structured.get("personal_info", {}),
    )
    db.add(resume)
    db.commit()
    db.refresh(resume)

    return {
        "id": resume.id,
        "filename": resume.filename,
        "skills": resume.skills,
        "experience": resume.experience,
        "projects": resume.projects,
        "education": resume.education,
        "personal_info": resume.personal_info,
    }


@router.get("/keywords")
async def get_resume_keywords(db: Session = Depends(get_db)):
    """Generate search keywords from the latest resume using AI."""
    resume = db.query(Resume).order_by(Resume.created_at.desc()).first()
    if not resume:
        return {"keywords": "", "skills": [], "positions": []}

    skills = resume.skills or []
    experience = resume.experience or []
    projects = resume.projects or []

    # Extract positions from work experience
    positions = []
    for exp in experience[:3]:
        pos = exp.get("position", "") if isinstance(exp, dict) else ""
        if pos and pos not in positions:
            positions.append(pos)

    # Use AI to generate smart search keywords
    try:
        from ..services.ai_service import chat as ai_chat
        import json as _json

        resume_summary = {
            "skills": skills,
            "positions": positions,
            "projects": [p.get("name", "") for p in projects[:3]] if projects else [],
        }

        prompt = f"""根据以下简历信息，生成5-8个适合在招聘网站（如BOSS直聘）搜索岗位的关键词。

简历摘要：
{_json.dumps(resume_summary, ensure_ascii=False, indent=2)}

规则：
1. 关键词必须是具体的职位名称，例如"Python开发工程师"而非"Python"，"数据分析师"而非"数据分析"
2. 优先从工作经历中提炼过往职位，并泛化为常见招聘职位名
3. 如果某技能没有直接的职位对应，可以组合为"XX工程师"或"XX开发"
4. 关键词用 / 分隔
5. 返回纯JSON：{{"keywords": "关键词1/关键词2/关键词3"}}

只返回JSON，不要其他内容。"""

        response = await ai_chat([
            {"role": "system", "content": "你是一个招聘搜索专家，精通招聘网站的搜索机制。你必须生成具体的职位名称作为关键词，而非技能名或工具名。只返回JSON。"},
            {"role": "user", "content": prompt},
        ])

        result = _json.loads(response.strip())
        keywords = result.get("keywords", "")
    except Exception:
        # Fallback: construct job-title-like keywords from positions + skills
        tool_keywords = {
            'microsoft office', 'wps', 'word', 'excel', 'ppt', 'powerpoint',
            'pr', 'ps', 'photoshop', '剪映', 'final cut', '美图秀秀',
            '美团商家后台', '抖音来客', '微信公众号',
        }
        useful_skills = [s for s in skills if s.lower() not in tool_keywords][:5]

        skill_to_title = {
            'python': 'Python开发工程师', 'java': 'Java开发工程师',
            'javascript': '前端开发工程师', 'typescript': '前端开发工程师',
            'react': 'React前端工程师', 'vue': 'Vue前端工程师',
            'golang': 'Go开发工程师', 'go': 'Go开发工程师',
            'rust': 'Rust开发工程师', 'c++': 'C++开发工程师',
            'c#': 'C#开发工程师', 'php': 'PHP开发工程师',
            'ruby': 'Ruby开发工程师', 'swift': 'iOS开发工程师',
            'kotlin': 'Android开发工程师', 'flutter': 'Flutter开发工程师',
            'sql': '数据库开发工程师', 'mysql': '数据库开发工程师',
            'postgresql': '数据库开发工程师', 'mongodb': '数据库开发工程师',
            'docker': 'DevOps工程师', 'kubernetes': 'DevOps工程师',
            'aws': '云计算工程师', 'azure': '云计算工程师',
            'tensorflow': '算法工程师', 'pytorch': '算法工程师',
            '机器学习': '机器学习工程师', '深度学习': '算法工程师',
            '数据分析': '数据分析师', '数据挖掘': '数据挖掘工程师',
            '产品': '产品经理', '运营': '运营专员',
            'ui': 'UI设计师', 'ux': 'UX设计师',
            '测试': '测试工程师', '自动化测试': '测试开发工程师',
            '运维': '运维工程师', '网络安全': '安全工程师',
        }
        skill_titles = []
        for s in useful_skills:
            title = skill_to_title.get(s.lower().strip(), '')
            if title and title not in skill_titles:
                skill_titles.append(title)

        all_parts = [p for p in positions if len(p) < 12] + [t for t in skill_titles if t not in positions]
        keywords = "/".join(all_parts[:6]) if all_parts else ""

    return {"keywords": keywords, "skills": skills, "positions": positions}


@router.get("/{resume_id}")
async def get_resume(resume_id: int, db: Session = Depends(get_db)):
    resume = db.query(Resume).filter(Resume.id == resume_id).first()
    if not resume:
        raise HTTPException(status_code=404, detail="简历不存在")
    return {
        "id": resume.id,
        "filename": resume.filename,
        "skills": resume.skills,
        "experience": resume.experience,
        "projects": resume.projects,
        "education": resume.education,
        "personal_info": resume.personal_info,
    }


@router.post("/update/{job_id}")
async def update_resume_for_job(
    job_id: int,
    resume_id: int,
    db: Session = Depends(get_db),
):
    resume = db.query(Resume).filter(Resume.id == resume_id).first()
    if not resume:
        raise HTTPException(status_code=404, detail="简历不存在")

    job = db.query(JobListing).filter(JobListing.id == job_id).first()
    if not job:
        raise HTTPException(status_code=404, detail="岗位不存在")

    # Generate optimized resume content via AI
    try:
        optimized_text = await update_resume(
            {
                "skills": resume.skills or [],
                "experience": resume.experience or [],
                "projects": resume.projects or [],
                "education": resume.education or [],
                "personal_info": resume.personal_info or {},
            },
            {
                "title": job.title,
                "company": job.company,
                "description": job.description,
                "requirements": job.requirements,
            },
        )
    except Exception:
        raise HTTPException(status_code=500, detail="AI生成失败，请稍后重试")

    # Build docx
    doc = Document()
    doc.styles["Normal"].font.name = "SimSun"
    doc.styles["Normal"].font.size = Pt(11)

    title_para = doc.add_paragraph()
    title_run = title_para.add_run(f"优化简历 - {job.title} @ {job.company}")
    title_run.bold = True
    title_run.font.size = Pt(14)

    doc.add_paragraph("")  # spacing

    for line in optimized_text.strip().split("\n"):
        line = line.strip()
        if not line:
            continue
        if line.startswith("**") or line.startswith("##"):
            # Section header
            clean = line.replace("**", "").replace("#", "").strip()
            p = doc.add_paragraph()
            run = p.add_run(clean)
            run.bold = True
            run.font.size = Pt(12)
        elif line.startswith("- ") or line.startswith("* "):
            # Bullet point
            doc.add_paragraph(line[2:], style="List Bullet")
        else:
            doc.add_paragraph(line)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    from urllib.parse import quote
    safe_name = f"resume_{job.id}.docx"
    encoded_name = quote(f"优化简历_{job.title}_{job.company}.docx")

    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": f"attachment; filename=\"{safe_name}\"; filename*=UTF-8''{encoded_name}",
        },
    )


@router.get("/")
async def list_resumes(db: Session = Depends(get_db)):
    resumes = db.query(Resume).order_by(Resume.created_at.desc()).all()
    return {
        "resumes": [
            {
                "id": r.id,
                "filename": r.filename,
                "skills": r.skills,
                "experience": r.experience,
                "projects": r.projects,
                "education": r.education,
                "personal_info": r.personal_info,
            }
            for r in resumes
        ]
    }
